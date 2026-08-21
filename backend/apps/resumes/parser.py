import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        extracted_text = []
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                extracted_text.append(text)
        return "\n\n".join(extracted_text)
    except Exception as e:
        logger.error(f"pypdf extraction error: {e}")
        raise ValueError(f"Could not parse PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX file bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        extracted_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    extracted_text.append(row_text)
        return "\n".join(extracted_text)
    except Exception as e:
        logger.error(f"python-docx extraction error: {e}")
        raise ValueError(f"Could not parse DOCX file: {str(e)}")

def parse_resume_document(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Determines file extension and extracts plain text.
    Returns (raw_text, file_type).
    """
    lower_name = filename.lower()
    if lower_name.endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
        return text, 'pdf'
    elif lower_name.endswith('.docx') or lower_name.endswith('.doc'):
        text = extract_text_from_docx(file_bytes)
        return text, 'docx'
    else:
        # Fallback to UTF-8 text decoding if plain text file
        try:
            text = file_bytes.decode('utf-8')
            return text, 'txt'
        except Exception:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX resume.")
